import csv
from docx import Document

csv_path='Data.csv'

with open(csv_path,'r') as csv_file:
    csv_reader=csv.DictReader(csv_file)
    print(csv_reader)
    for row in csv_reader:
        doc = Document('Template.docx')

        for variable, value in row.items():
            search_string = f'{{{variable}}}'
            print(search_string)
            for paragraph in doc.paragraphs:
                print(paragraph.text)
                if search_string in paragraph.text:
                    paragraph.text = paragraph.text.replace(search_string, str(value))

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if search_string in cell.text:
                            cell.text = cell.text.replace(search_string, str(value))

        output_file_path = f'output/output_{row["COMPANY_NAME"]}_proposal.docx'
        doc.save(output_file_path)

print("Documents populated and saved.")